import csv
import os
import docx
from docx import Document
from xlsx2csv import Xlsx2csv
from pathlib import Path
from tkinter import  filedialog
from collections import defaultdict
from typing import List, Tuple

from logo import createLogoImage
from student import Student  
from letter import LetterWriter
from language_map import LANGUAGE_MAP

def select_file():
    return filedialog.askopenfilename(title="Select Scoresheet File", filetypes=[("Excel files", "*.xlsx"), ("CSV files", "*.csv")])

def open_file(filename):
    if not filename:
        return

    data = []
    if filename.endswith('.xlsx'):
        temp_csv = filename.replace('.xlsx', '.csv')
        Xlsx2csv(filename).convert(temp_csv)
        filename = temp_csv

    with open(filename, 'r') as file:
        reader = csv.reader(file)
        headers = next(reader)
        for row in reader:
            record = {}
            for i, value in enumerate(row):
                record[headers[i]] = value
            data.append(record)
    return data

def setup_document(teacher_email, log_label): 
    log_label["text"] = "Creating Word Document..."
    log_label.update()
    # Create a document
    doc = Document()

    # Add header with logo on the left and details on the right
    header = doc.sections[0].header
    header_table = header.add_table(rows=1, cols=2, width=doc.sections[0].page_width)
    header_table.columns[0].width = docx.shared.Inches(1.5)
    header_table.columns[1].width = doc.sections[0].page_width - docx.shared.Inches(1.5)

    # Add the school logo to the left cell
    logo_cell = header_table.rows[0].cells[0]
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
    logo_run = logo_paragraph.add_run()
    
    logo_run.add_picture(createLogoImage(), width=docx.shared.Inches(1.0), height=docx.shared.Inches(1.0))

    # Add school details to the right cell
    details_cell = header_table.rows[0].cells[1]
    details_paragraph = details_cell.paragraphs[0]
    details_paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT

    details_paragraph.add_run("Western International High School\n1500 Scotten, Detroit, MI 48209\n{}\n".format(teacher_email))
    details_paragraph.paragraph_format.space_after = docx.shared.Pt(6)

    return doc

def parse_students(filename, log_label, settings, filters=[]):
    log_label["text"] = "Processing file..."
    log_label.update()
    data = open_file(filename)
    students = []
    for entry in data:
        first_column = list(entry.keys())[0]
        name = ' '.join(reversed(entry[first_column].split(', ')))
        student = Student(entry, settings.class_name, settings.custom_message)

        if all(filter.apply(student) for filter in filters):
            students.append(student)
    return students
    
def generate_report(log_label, settings, filters=[]):
    filename = select_file()
    students = parse_students(filename, log_label, settings, filters)
    doc = generate_report_for_language(students, log_label, settings, settings.language)
    save_document(doc, log_label, settings)

def generate_report_for_language(students, log_label, settings, language):
    doc = setup_document(settings.teacher_email, log_label)
    letterWriter = LetterWriter(settings.teacher_name, settings.teacher_email, language, settings.custom_message)

    log_label["text"] = "Translating text... This may take up to a minute."
    log_label.update()
    for student in students:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = 0
        run = p.add_run(letterWriter.generateLetter(student))
        run.font.name = 'Arial'
        run.font.size = docx.shared.Pt(12)
        doc.add_page_break()
    return doc

def save_document(doc, log_label, settings):
    # Save the document
    log_label["text"] = "Saving document..."
    log_label.update()

    output_filename = "{}.docx".format(settings.class_name.replace(" ", "_"))
    output_path = Path.home() / "Downloads" / output_filename
    doc.save(output_path)

    log_label["text"] = "Completed. File saved to downloads as {}.".format(output_filename)
    log_label.update()
    os.startfile(output_path)

def generate_report_for_selected(log_label, settings, student_language_pairs: List[Tuple[Student, str]]):
    log_label["text"] = "Grouping students by language..."
    log_label.update()

    language_grouped_students = defaultdict(list)
    for student, language in student_language_pairs:
        language_grouped_students[language].append(student)

    print(f"Grouped students by language: {language_grouped_students.keys()}")
    doc = setup_document(settings.teacher_email, log_label)

    for language, students in language_grouped_students.items():
        log_label["text"] = f"Translating and writing {LANGUAGE_MAP[language]} letters..."
        log_label.update()
        print(f"Generating report for {language} with {len(students)} students")
        generate_report_for_language(students, log_label, settings, language)

    save_document(doc, log_label, settings)